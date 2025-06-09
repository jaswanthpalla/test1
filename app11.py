import streamlit as st
import re
import spacy
import subprocess
from langchain import PromptTemplate, LLMChain
from langchain_huggingface import HuggingFaceEndpoint
from docx import Document
from dotenv import load_dotenv
import os

# Load environment variables (for local development)
load_dotenv()

# Get API key from environment (works both locally and on Streamlit Cloud)
HF_API_KEY = os.getenv("HF_API_KEY")

# Verify API key is loaded
if not HF_API_KEY:
    st.error("HF_API_KEY not found. Set it in .env (locally) or Streamlit Secrets (on deployment).")
    st.stop()

# Download SpaCy model if not already installed
nlp = spacy.load("en_core_web_sm")

# Initialize HuggingFace LLM (Mistral model) using HuggingFaceEndpoint
llm = HuggingFaceEndpoint(
    endpoint_url="https://api-inference.huggingface.co/models/mistralai/Mixtral-8x7B-Instruct-v0.1",
    huggingfacehub_api_token=HF_API_KEY,
    temperature=0.7,
    max_new_tokens=500
)

# Define prompt templates for LangChain
summary_prompt = PromptTemplate(
    input_variables=["transcript"],
    template="Summarize the following meeting transcript in 100 words or less:\n\n{transcript}"
)

discussion_points_prompt = PromptTemplate(
    input_variables=["transcript"],
    template="List the key discussion points from the following meeting transcript as bullet points:\n\n{transcript}"
)

action_items_prompt = PromptTemplate(
    input_variables=["transcript"],
    template="Extract action items from the following meeting transcript. For each action item, specify the task and, if possible, the assignee. Format as:\n- Task: [task description], Assignee: [name or None]\n\n{transcript}"
)

# Create LLM chains
summary_chain = LLMChain(llm=llm, prompt=summary_prompt)
discussion_points_chain = LLMChain(llm=llm, prompt=discussion_points_prompt)
action_items_chain = LLMChain(llm=llm, prompt=action_items_prompt)

# Function to clean transcript
def clean_transcript(transcript):
    # Remove timestamps (e.g., 00:01:23 --> or [00:01])
    transcript = re.sub(r'\d{2}:\d{2}:\d{2}\s*-->\s*\d{2}:\d{2}:\d{2}', '', transcript)
    transcript = re.sub(r'\[\d{2}:\d{2}\]', '', transcript)
    # Remove extra whitespace and normalize
    transcript = re.sub(r'\s+', ' ', transcript).strip()
    return transcript

# Function to extract assignees using SpaCy
def extract_assignees(transcript):
    doc = nlp(transcript)
    names = [ent.text for ent in doc.ents if ent.label_ == "PERSON"]
    return list(set(names))  # Remove duplicates

# Function to refine action items with assignees
def refine_action_items(action_items, names):
    refined = []
    for item in action_items.split('\n'):
        if item.strip():
            assignee = "None"
            for name in names:
                if name.lower() in item.lower():
                    assignee = name
                    break
            refined.append(f"{item}, Assignee: {assignee}")
    return refined

# Function to create .docx file
def create_docx(summary, discussion_points, action_items):
    doc = Document()
    doc.add_heading("Meeting Summary", 0)
    
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(summary)
    
    doc.add_heading("Key Discussion Points", level=1)
    doc.add_paragraph(discussion_points)
    
    doc.add_heading("Action Items", level=1)
    for item in action_items:
        doc.add_paragraph(item, style="ListBullet")
    
    doc_path = "meeting_summary.docx"
    doc.save(doc_path)
    return doc_path

# Streamlit app
st.title("Meeting Notes Summarizer")

# File uploader
uploaded_file = st.file_uploader("Upload meeting transcript (.txt)", type="txt")

if uploaded_file:
    # Read and clean transcript
    transcript = uploaded_file.read().decode("utf-8")
    cleaned_transcript = clean_transcript(transcript)
    
    # Display cleaned transcript
    st.subheader("Cleaned Transcript")
    st.text_area("Transcript", cleaned_transcript, height=200)
    
    if st.button("Process Transcript"):
        with st.spinner("Processing..."):
            # Generate summary
            summary = summary_chain.run(transcript=cleaned_transcript)
            
            # Extract discussion points
            discussion_points = discussion_points_chain.run(transcript=cleaned_transcript)
            
            # Extract action items
            action_items = action_items_chain.run(transcript=cleaned_transcript)
            
            # Extract assignees using SpaCy
            names = extract_assignees(cleaned_transcript)
            refined_action_items = refine_action_items(action_items, names)
            
            # Display results
            st.subheader("Summary")
            st.write(summary)
            
            st.subheader("Key Discussion Points")
            st.write(discussion_points)
            
            st.subheader("Action Items")
            for item in refined_action_items:
                st.write(item)
            
            # Create and provide download link for .docx
            doc_path = create_docx(summary, discussion_points, '\n'.join(refined_action_items))
            with open(doc_path, "rb") as file:
                st.download_button(
                    label="Download Summary as .docx",
                    data=file,
                    file_name="meeting_summary.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# Instructions for running (for local testing)
st.markdown("""
### How to Run Locally
1. Install dependencies: `pip install -r requirements.txt`
2. Download SpaCy model: `python -m spacy download en_core_web_sm`
3. Create a `.env` file with `HF_API_KEY=your_api_token`
4. Run: `streamlit run meeting_summarizer.py`
5. Upload a .txt transcript and click 'Process Transcript'.
""")
